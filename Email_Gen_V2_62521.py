# -*- coding: utf-8 -*-
"""
This code was developed to automatically generate emails (word docs) for SSSP by pulling superintendent and school information from excel file for School contact list


Created on Thu Oct  3 09:55:49 2019
Washington Department of Natural Resources Geological Survey
Alex Wernle 
alex.wernle@dnr.wa.gov

Last updated 10/22/2019

@author: awer490
"""
##############################################################################
#Defined your PWD
pwd= 'C:/Users/awer490/Desktop/AW_SSSP/Python_ish/Email_Py_Script/' 

# Import modules
import pandas as pd
import docx
from docx import Document
from docx.shared import RGBColor
from docx.enum.dml import MSO_THEME_COLOR_INDEX
from docx.shared import Pt

# read Excel documents (CHANGE THIS STUFF FOR NEW SHEETS)

# open excel doc with all schools=> must be located in the same directory
df = pd.read_excel("All.PLANNING.phase2.xlsx", sheet_name= 'participants')
# open shortened excel doc with just districts
df2 = pd.read_excel('All.PLANNING.phase2_districts.xlsx')
# open excel doc with all POC info
df3 = pd.read_excel("School_POC.phase2.xlsx", sheet_name= 'Revised_poc')
# open excel doc with building info for table at end of email 
df4 = pd.read_excel("Copy of 210602 Ph 1 and 2 Prioritization Spreadsheet to DNR.xlsx", sheet_name= 'DNR School Bldg Prioritization')


#Define your excel columns (you many not need all of these depending on email)
School = df['SiteName'] 
Districts = df['DistrictName']
Districts2 = df2['DistrictName']
Districts3 = df3['DistrictName']
POC_SD = df3['POCSDname']
POC_1 = df3['POCSD_email']
POC_2 = df3['Principal_Email']
POC_3 = df3['other3_email']
POC_4 = df3['other1_email']
POC_5 = df3['other2_email']
POC_6 = df3['other3_email.1']
POC_7 = df3['other4_email']
Districts4 = df4['School District']
Schools4 = df4['SiteName']
Building = df4['Building']
ICOS = df4['ICOS']

#convert to list for enumeration
Districts3 = Districts3.tolist()

#create array to append to for all POCs
POCS = []


# Ignore this:
##############################################################################
#Define the hyperlink function,  
#Reference: https://stackoverflow.com/questions/47666642/adding-an-hyperlink-
#in-msword-by-using-python-docx
def add_hyperlink(paragraph, text, url):
    # Gets access to the document.xml.rels file and gets a new relation id 
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, 
                          is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a w:r element and a new w:rPr element
    new_run = docx.oxml.shared.OxmlElement('w:r')
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Join all the xml elements together add add the required text to the w:r
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    # Create a new Run object and add the hyperlink into it
    r = paragraph.add_run ()
    r._r.append (hyperlink)

    # A workaround for the lack of a hyperlink style (doesn't turn purple)
    # Delete this if using a template that has the hyperlink style in it
    r.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
    r.font.underline = True 

    return hyperlink
##############################################################################

#Iterate through district column
for dirname in Districts2:
    
    #Create document for email
    document = Document()
    document.save (dirname)
    
    #Append correct POCs to array to be printed at the bottom of email
    for index,y in enumerate(Districts3):
        if dirname == y:  
            pos = index
            if type(POC_1[pos]) is str: 
                POCS.append(POC_1[pos]+";") 
            if type(POC_2[pos]) is str:
                POCS.append(POC_2[pos]+";")                
            if type(POC_3[pos]) is str:
                POCS.append(POC_3[pos]+";")
            if type(POC_4[pos]) is str:
                POCS.append(POC_4[pos]+";")
            if type(POC_5[pos]) is str:
                POCS.append(POC_5[pos]+";")
            if type(POC_6[pos]) is str:
                POCS.append(POC_6[pos]+";")                 
            if type(POC_7[pos]) is str:                
                POCS.append(POC_7[pos]+";")
            
    #removes duplicates from POCs        
    POCS = list(dict.fromkeys(POCS))
    #remves unknowns from POCS
    elem = "unknown;"
    if elem in POCS:
        POCS.remove(elem)
        
    #define Superintendent and Schools for the district
    Superintendent = ((POC_SD[pos]))
    # Split string to only keep last name of supertintendent
    Out = (Superintendent.split()[-1])
    # Define school info at the row for current school(dirname)
    Schools = ( (School[Districts == dirname]) ) 
    # convert to series for iteration
    Schools = pd.Series(Schools)
    
    # Loop through school series and split string to get rid of identifier
    Schools_array = []
    for school in Schools:
        str_split=' '.join(school.split()[0:])#split string
        Schools_array.append(str_split)


    #write to word file (CHANGE TEXT HERE)####################################
    #create first paragraph
    doc_para= document.add_paragraph('Dear Superintendent ')  
    
    #Define paragraph style!
    style = document.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    doc_para.style=document.styles['Normal']
  
    # Add superintendent info from above
    doc_para.add_run(Out)
    doc_para.add_run(", Principals, and School Staff,")

    doc_para= document.add_paragraph("""The following school(s) in your district were assessed in Phase 2 of the School Seismic Safety Project, for a list of the buildings assessed at these schools see the table at the end of this ridiculously long email:""")
    
    #create bulleted list for each school in the school array
    for schools in Schools_array:
        doc_para2= document.add_paragraph(schools)
        doc_para2.style = 'List Bullet'

    #add another paragraph
    doc_para2= document.add_paragraph("""    	I am writing to provide the final report on Phase 2 of the statewide School Seismic Safety Project (SSSP), as appropriated in the 2019–2021 capital budget, led by the Washington State Department of Natural Resources (DNR) Washington Geological Survey (WGS) with significant contributions from structural engineering contractors led by Reid Middleton, Inc., and the Office of Superintendent of Public Instruction (OSPI). The SSSP school assessments are based on local geology and the engineering and construction of the buildings studied. The attached report summarizes the seismic risk at 561 school buildings (274 schools at 245 campuses) across the state and is the culmination of two biennia of work; Phase 1, which was funded by the 2017–2019 capital budget, and Phase 2, which was funded by the 2019–2021 capital budget. This report presents the results of the 339 buildings studied in Phase 2 (2019–2021 biennium appropriation), with some high level conclusions from Phases 1 and 2 combined.""")
    
    
    sent = doc_para2.add_run("""
Project Overview""")
    sent.font.name = 'Calibri Light'
    sent.font.color.rgb = RGBColor(0x42, 0x24, 0xE9)
    sent.font.size = Pt(13)
    
    doc_para2.add_run("""
    	The project involves both geological and engineering assessments at each school. Geologists collected seismic data to measure how local soils amplify earthquake shaking at school campuses, usually on playing fields. This seismic data greatly improves estimates of potential ground shaking by more accurately evaluating site-specific soil conditions under the school buildings. In addition to this, a group of licensed professional structural engineers collected building data at the schools. The structural and nonstructural adequacy of the school buildings were evaluated and safety ratings and damage estimates for these buildings were developed. Combined, these assessments provide a detailed view of how earthquake shaking might affect each school. A selection of high-risk buildings were studied in more detail to determine what a seismic retrofit design would look like and estimate how much it would cost to complete that upgrade. These are called ‘concept-level seismic upgrade designs’. For schools that received a concept-level design a separate email with more information regarding that process will be sent.""")
    
    sent= doc_para2.add_run("""
    Key Results""")
    sent.font.name = 'Calibri Light'
    sent.font.color.rgb = RGBColor(0x42, 0x24, 0xE9)
    sent.font.size = Pt(13)

    
    doc_para2.add_run("""
    •	Washington State has many older school buildings built prior to the adoption of modern seismic safety codes. Older and more vulnerable construction types are more susceptible to earthquake damage and have a greater percentage of seismically noncompliant structural and non-structural components.""")
    doc_para2.add_run("""
    •	Unreinforced masonry buildings constructed before the 1940s and non-ductile concrete buildings (without seismic upgrades) constructed before the mid-1970s located in high seismic hazard areas are especially vulnerable to collapse during earthquakes. The risks of these buildings should be mitigated as soon as practical.""")
    doc_para2.add_run("""
    •	In total, 67 school buildings on 30 school campuses that were assessed in Phase 1 and Phase 2 are located within tsunami inundation zones. These schools serve more than 10,000 students. Tsunami loads and impacts were not considered in the geologic or engineering assessments. For schools to be safe from a tsunami, they would need to be moved from the tsunami inundation zone or designed to withstand tsunami loads with options for vertical evacuation.  """)
    doc_para2.add_run("""
    •	The concept-level seismic upgrade design results indicate that for many buildings, the cost to seismically upgrade the structure will cost less than the costs to repair major damage following an earthquake, or significantly less than the cost to replace an irreparably damaged building. For less vulnerable structures, especially structures in low seismicity areas, however, it may not be financially worth implementing seismic upgrades. """)
    doc_para2.add_run("""
    •	A range of cost estimates were developed for each of the select buildings that received a concept-level design and estimated costs to retrofit. Phase 1 concept-level design building cost estimates ranged from $63K to $5.01M, and were for buildings geographically spread across the state. These cost estimates were for construction costs only and did not include an allowance for project soft costs.  Phase 2 concept level design building cost estimates ranged from $1.24M to $15.26M. Cost estimate methods for Phase 2 were improved from Phase 1 and now include projected soft costs. Phase 1 concept design schools were selected to represent a variety of building construction types and vintages in different seismic hazard areas. Alternatively, Phase 2 concept design schools were selected based on available information to be some of the highest risk buildings based on seismic hazard and engineering design.  """)
    doc_para2.add_run("""
    •	A significant portion of the structural upgrade costs are due to the fact that the seismic upgrades take place in existing buildings with existing finishes and existing nonstructural components. The costs to temporarily remove and replace the architectural, mechanical, electrical, and plumbing equipment is significant. If the costs associated with the architectural, mechanical, electrical, plumbing, and fire protection elements were deleted from the cost estimates, the average seismic upgrade cost sees a 70 percent reduction. Significant savings can be realized by combining seismic upgrades with other types of work, such as re-roofing projects or school modernizations.  """)
    doc_para2.add_run("""
    •	Phase 1 and 2 school buildings were ranked to prioritize buildings for seismic retrofit by relative risk. Of the 561 buildings studied, 63 percent were high or very high priority, 18 percent were moderate priority, and 19 percent were lower priority.  """)
    
    sent = doc_para2.add_run("""
Next Steps""")
    sent.font.name = 'Calibri Light'
    sent.font.color.rgb = RGBColor(0x42, 0x24, 0xE9)
    sent.font.size = Pt(13)

    doc_para2.add_run(""" 
    	The School Seismic Safety Project (Phases 1 and 2) has been an important opportunity to study and evaluate school buildings across the state and has demonstrated the need for dedicated funding for seismic retrofits. Following the Phase 1 report and project, the Legislature funded OSPI $13 million in 2019 and $40 million in the 2021–2023 biennium for the School Seismic Safety Retrofit Program (SSSRP). This program is the first of its kind in Washington and is a critical step in repairing the most vulnerable schools. For more information about this Program please contact Scott Black: """) 

    add_hyperlink(doc_para2, """Scott.black@k12.wa.us""", "mailto:Scott.black@k12.wa.us")

    doc_para2.add_run(""" 
    	The State of Washington has adopted the 2018 International Existing Building Code as its building standard for existing buildings. Per this building code, a school district is under no obligation to upgrade its school buildings to the suggested recommendations unless there is a change in use or occupancy, an addition, or a significant alteration made to the existing structure. 
        The results in this report are therefore informational, with the intent of helping districts, schools, parents, legislators, OSPI, and the public better understand the seismic risk at Washington school campuses. Public schools will need financial support to make the structural improvements outlined here.""")
    sent = doc_para2.add_run("""
Where to Find Reports for your District""")
    sent.font.name = 'Calibri Light'
    sent.font.color.rgb = RGBColor(0x42, 0x24, 0xE9)
    sent.font.size = Pt(13)
    doc_para2.add_run(""" 
        I have attached the final report for Phase 2 of the project, which contains summary information of the process, major findings, recommendations, and links to individual reports and prioritizations for all schools assessed. School building and district specific information is available from links in Appendix A and B of the report and also on our website: """)
    add_hyperlink(doc_para2, 'School Seismic Safety Project webpage', "https://www.dnr.wa.gov/school-seismic-safety") 
    doc_para2.add_run(""". To download the individual engineering and geology reports for your district use this map: """)
    add_hyperlink(doc_para2, 'WA School Seismic Safety Project Assessments', "https://experience.arcgis.com/experience/3451c6fb71f347b2b9d302cf77fae3d2/?data_id=dataSource_1-schools_and_reports_shapefile_2961%3A2 ")
    doc_para2.add_run(""" 
        If you have any questions about the geologic or engineering information provided for your school(s) please do not hesitate to reach out to me with questions. 

Sincerely,
""")
                   


    #Add Table here###########################################################
    #Define array to establish table dimensions
    Schools_array = []
    for district in Districts4:
        if district == dirname:
            Schools_array.append(district)
    
    #Define other arrays for table
    district_names=(Districts4[Districts4==dirname])
    school_names=(Schools4[Districts4==dirname])
    Bldg_names=(Building[Districts4==dirname])
    ICOS_numbers=((ICOS[Districts4==dirname])).astype(str)
            
            
    row_count = (len(Schools_array))+1
    # create table and define size
    table = document.add_table(rows=1, cols= 4)
    table.autofit= True
    table.style = 'TableGrid'
    
    # define column headings
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'School District'
    hdr_cells[1].text = 'School Name'
    hdr_cells[2].text = 'Building Name'
    hdr_cells[3].text = 'ICOS Number'

    #Define information in rows
    for ind in range(0, len(Schools_array)): 
       row_cells = table.add_row().cells
       row_cells[0].text = district_names.iloc[ind]
       row_cells[1].text = school_names.iloc[ind]
       row_cells[2].text = Bldg_names.iloc[ind]
       row_cells[3].text = ICOS_numbers.iloc[ind] 
    ##########################################################################

    # Add a final paragraph with POCS printed out
    doc_para3= document.add_paragraph(
        POCS)
    
    #End file and save
    document.save (dirname)
    POCS.clear()
##############################################################################   
