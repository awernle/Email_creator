# -*- coding: utf-8 -*-
"""
Created on Thu Oct  3 09:55:49 2019
Washington Department of Natural Resources Geological Survey
Alex Wernle 
alex.wernle@dnr.wa.gov

Last updated 10/22/2019

@author: awer490
"""
#Alex Wernle 10/3/19
# the purpose of this script is to automatically generate emails( word docs) for SSSP by pulling key words from excel file for School contact list

 #Present working directory
pwd= 'C:/Users/awer490/Desktop/AW_SSSP/Python_ish/Email_Py_Script/'

import pandas as pd
import docx
from docx import Document
from docx.enum.dml import MSO_THEME_COLOR_INDEX
from docx.shared import Pt

# open excel doc with all schools
df = pd.read_excel('Schools_Contact_List_Short-Schools.xls') 
# open shortened excel doc with just districts
df2 = pd.read_excel('Schools_Contact_List_Short-Districts.xls')


#Define your excel columns
School = df['SiteName'] 
Districts = df2['DistrictNa'] 
Districts2 = df['DistrictNa'] 
POC_SD = df2['POC SD Level Contact']


#Define the hyperlink function################################################
def add_hyperlink(paragraph, text, url):
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a w:r element and a new w:rPr element
    new_run = docx.oxml.shared.OxmlElement('w:r')
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    # Create a new Run object and add the hyperlink into it
    r = paragraph.add_run ()
    r._r.append (hyperlink)

    # A workaround for the lack of a hyperlink style (doesn't go purple after using the link)
    # Delete this if using a template that has the hyperlink style in it
    r.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
    r.font.underline = True

    return hyperlink
##############################################################################

#Loop through district column
for dirname in Districts:
    
    #Create document for email
    document = Document()
    document.save (dirname)

    #define Superintendent and Schools for the district
    Superintendent = ((POC_SD[Districts == dirname]).to_string())
    Out = ' '.join(Superintendent.split()[2:])
    Schools = ( (School[Districts2 == dirname])) 
    Schools = pd.Series(Schools)
    
    # Loop through school series and split string to get rid of object identifier
    Schools_array = []
    for school in Schools:
        str_split=' '.join(school.split()[0:])#split string
        Schools_array.append(str_split)

    #write to word file (CHANGE TEXT HERE)####################################
    doc_para= document.add_paragraph('Dear Superintendent ')  #create first paragraph
    
    #Define paragraph style
    style = document.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    doc_para.style=document.styles['Normal']
  
    
    doc_para.add_run(Out)# Add superintendent info from above
    doc_para.add_run(",")

    doc_para= document.add_paragraph("""My name is Travis West, I am the lead geophysicist on the Washington State School Seismic Safety Project (SSSP), an initiative led by the Department of Natural Resources’ (DNR) Washington Geological Survey (WGS). The SSSP is a statewide project to evaluate how Washington public school buildings could be affected by earthquake shaking. The Washington State Legislature has provided  capital funding for this initiative in the 2019–2021 biennium.
I am contacting you in partnership with the Office of Superintendent of Public Instruction (OSPI) because a school(s) in your district has been identified for possible participation in Phase 2 of the SSSP. We would like to take this opportunity to provide you with more information about the project and to describe what your participation would entail. These are the school(s) we have preliminarily selected:""")
    
    
    #create bulleted list for each school in the school array
    for schools in Schools_array:
        doc_para2= document.add_paragraph(schools)
        doc_para2.style = 'List Bullet'

    
    doc_para2= document.add_paragraph("""Phase 1 of the SSSP was recently completed for 222 school buildings throughout Washington State. The assessment consists of two parts: (1) a seismic site class evaluation of soil conditions, and (2) detailed inspections of the exterior and interior of the buildings by a licensed structural engineer. These parts require two site visits by two separate teams. Results for these buildings were incorporated into OSPI’s Information and Condition of Schools (ICOS) database and distributed to each school and district. An overview of the project and the Phase I final reports can be accessed at the WGS """)
    
    add_hyperlink(doc_para2, 'School Seismic Safety Project webpage', "https://www.dnr.wa.gov/programs-and-services/geology/geologic-hazards/earthquakes-and-faults/school-seismic-safety") 

    doc_para2.add_run(""". Phase 2 follows the same procedure as Phase 1.

""")
    #! ------- Change me if a link is provided to upload building plans -------
    doc_para2.add_run('Building Plans:').bold = True

    doc_para2.add_run(""" In order to determine if your school(s) would be a good candidate for a seismic safety assessment, we need to examine the school’s blueprints. Having the building’s blueprints enhances the effectiveness of the detailed seismic inspection performed by the structural engineers. We are therefore giving higher priority to schools that have access to their building’s structural drawings/plans (blueprints). If you would like your school(s) to participate in this project we would appreciate if you could please reach out to your facilities manager to locate these blueprints and get back to us on whether they are available either digitally or in hard copy. If a digital version of the blueprints already exists, please inform us and we will follow up with a link for it's submission shortly. Otherwise, we will have a member of our team work with you to view them and see if we can scan them at a later time. 

Once we hear back from you about the availability of building plans, we will make a final selection of which buildings we will assess during Phase 2. If your building(s) is not selected for this phase of the project, we may reach out to you again at a later date for a potential Phase 3.

""") 

    doc_para2.add_run('School Scheduling:').bold = True
    
    doc_para2.add_run(""" If your school(s) is selected, we will schedule a date for our geophysics field team to visit your schools sometime this fall or next spring. Typically, our team requires a few hours to collect seismic data outside the building (usually in the ball field or a nearby open space) using methods that will not damage the property. As long as we have access to the field site, it is not necessary for a school representative to be present for the geophysical testing. However, we are pleased to accommodate any requests you may have. The engineering firm will contact you about scheduling a separate visit to evaluate the buildings. 

""")

    doc_para2.add_run('Educational Outreach:').bold = True

    doc_para2.add_run(""" Schools might view our site visits as opportunities for their students to learn about earth science and engineering. We would be happy to demonstrate our field methods of “listening” to seismic waves and briefly discuss their use with your students. If this demonstration is of interest to you, we can discuss coordinating this activity with our field testing. 

Please contact us at your convenience to discuss participation and scheduling. We are also happy to answer any other questions you may have about the project. You may respond to this e-mail directly or contact Scott Black (OSPI) or Corina Forson (WGS) using the information provided below.

Sincerely,""")


                     
    doc_para2.add_run("""

Loyd “Travis” West""").bold = True
    doc_para2.add_run("""
Washington Geological Survey
1111 Washington St SE, MS 47007
Olympia, WA 98504-7007
Office: 360-902-1481
Cell: 360-764-0347""")
    add_hyperlink(doc_para2, """
Travis.West@dnr.wa.gov""", "mailto:Travis.West@dnr.wa.gov")

    doc_para2.add_run("""

Alex Wernlé""").bold = True
    doc_para2.add_run("""
SSSP Assistant Geophysicist
Washington Geological Survey
360-902-2174 (office)""")
    add_hyperlink(doc_para2, """
Alex.Wernle@dnr.wa.gov""", "mailto:Alex.Wernle@dnr.wa.gov") 


    doc_para2.add_run("""

Scott Black""").bold = True
    doc_para2.add_run("""
Program Development Manager
Office of the Superintendent of Public Instruction
(360) 725-6268 (office)""")
    add_hyperlink(doc_para2, """
Scott.black@k12.wa.us""", "mailto:Scott.black@k12.wa.us") 


    doc_para2.add_run("""

Corina Forson""").bold = True
    doc_para2.add_run("""
Chief Hazards Geologist
Washington Geological Survey
(360) 902-1455 (office)
(360) 791-0647 (cell)""")
    add_hyperlink(doc_para2, """
Corina.forson@dnr.wa.gov""", "mailto:Corina.forson@dnr.wa.gov") 

    #End file write
    document.save (dirname)
    
##############################################################################   
